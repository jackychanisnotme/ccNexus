from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/pc/Documents/New project 2")
OUT = ROOT / "output" / "AINexus-6.7.2-介绍和使用手册.docx"
ASSET = ROOT / "output" / "ainexus_672_manual_assets"
VERSION = "6.7.2"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(28, 36, 50)
MUTED = RGBColor(91, 102, 120)
BORDER = "D0D7DE"
HEADER_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"


def set_run_font(run, size=None, bold=None, color=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        tc_pr.append(node)
    node.set(qn("w:fill"), fill)


def cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in [("top", top), ("bottom", bottom), ("start", start), ("end", end)]:
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def table_borders(table, color=BORDER, size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def configure(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    for style_name in ["List Bullet", "List Number"]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.38)
        style.paragraph_format.first_line_indent = Inches(-0.19)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.text = ""
    r = header.add_run(f"AINexus {VERSION} 最新版介绍和使用手册")
    set_run_font(r, size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = footer.add_run(f"AINexus v{VERSION} 操作参考")
    set_run_font(fr, size=9, color=MUTED)


def para(doc, text="", style=None, size=None, bold=False, color=None, before=0, after=6):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold, color=color)
    return p


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        r = p.add_run(item)
        set_run_font(r)


def numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        r = p.add_run(item)
        set_run_font(r)


def add_table(doc, headers, rows, widths=None, font_size=8.8):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table_borders(table)
    repeat_header(table.rows[0])
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade(cell, HEADER_FILL)
        cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(str(text))
        set_run_font(r, size=font_size, bold=True, color=INK)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cell = cells[i]
            cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            r = p.add_run(str(text))
            set_run_font(r, size=font_size)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    para(doc, "", after=4)
    return table


def callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.3)
    table_borders(table, color="B8C7D9", size="6")
    cell = table.cell(0, 0)
    shade(cell, CALLOUT_FILL)
    cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, bold=True, color=DARK_BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.2
    r2 = p2.add_run(body)
    set_run_font(r2, size=10)
    para(doc, "", after=4)


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_run_font(r, size=9, color=MUTED, bold=True)


def image(doc, filename, text, width=6.25):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(ASSET / filename), width=Inches(width))
    caption(doc, text)


def cover(doc):
    para(doc, "AINexus", size=30, bold=True, color=RGBColor(0, 0, 0), after=2)
    para(doc, f"v{VERSION} 最新版介绍和使用手册", size=18, color=DARK_BLUE, after=12)
    para(doc, "桌面客户端模式 / 服务器模式 / 授权服务器后台 / Token Pool / 远程维护", size=12, color=MUTED, after=20)
    icon = ROOT / "cmd" / "desktop" / "build" / "appicon.png"
    if icon.exists():
        p = doc.add_paragraph()
        p.add_run().add_picture(str(icon), width=Inches(1.1))
    callout(
        doc,
        "适用范围",
        "本手册适用于 AINexus 6.7.2，覆盖桌面客户端模式、服务器模式、授权服务器后台、Token Pool、数据同步、Agent Provider、远程端点维护和常见接入流程。界面字段和按钮说明以 6.7.2 功能为准。",
    )
    add_table(
        doc,
        ["项目", "内容"],
        [
            ["产品定位", "Codex / Claude Code / OpenClaw / Hermes Agent 的本地 API Provider、Token Pool 与 Agent 管理中枢"],
            ["默认代理地址", "桌面客户端：http://127.0.0.1:3000；服务器 Web UI：http://127.0.0.1:3000/ui/"],
            ["授权方式", "在线卡密激活；Ed25519 票据；最近一次在线校验成功后可离线宽限 30 天"],
            ["6.7.2 重点", "远程端点维护、端点错误遥测、Codex Token Pool 额度/重置次数/凭证级用量、Agent Provider 修复"],
            ["适用对象", "桌面客户端用户、服务器部署者、授权后台管理员和售后维护人员"],
        ],
        widths=[1.35, 4.95],
        font_size=9.3,
    )
    doc.add_page_break()


def overview(doc):
    doc.add_heading("1. 产品概览与 6.7.2 重点", level=1)
    para(doc, "AINexus 是面向 AI 编程工具和兼容客户端的智能 API 端点轮换代理。它统一管理多个上游 API、API Key、订阅 Token、Codex 登录凭据、授权状态、备份、统计和 Agent Provider 配置，并在请求失败时自动切换到可用端点。")
    bullets(
        doc,
        [
            "多端点轮换与自动故障转移：端点顺序、启用状态、冷却状态、并发限制和故障分类共同影响请求选择。",
            "协议转换：支持 Claude、OpenAI Chat、OpenAI Responses、Gemini、DeepSeek、Kimi、Poe 等转换器。",
            "Token Pool：支持普通 API Token Pool、Codex Token Pool 和实验性 Claude OAuth Token Pool。",
            "6.7.2 运维能力：授权后台可查看远程端点快照、下发端点维护命令、查看端点错误遥测，并保持密钥脱敏与最小上传。",
            "桌面与服务器双模式：桌面模式适合本机 GUI 与启动器；服务器模式适合 NAS、Docker 或局域网共享 API Provider。",
        ],
    )
    add_table(
        doc,
        ["运行模式", "入口", "主要界面", "适合场景"],
        [
            ["桌面客户端", "cmd/desktop", "单页主界面 + 设置、端点、Token Pool、备份、启动器、Agent 弹窗", "个人本机使用、托盘常驻、快速接入 Codex/Claude"],
            ["服务器模式", "cmd/server", "浏览器 Web UI：仪表板、端点、统计、测试、设置", "服务器、NAS、Docker、团队内共享 API Provider"],
            ["授权服务器后台", "cmd/license-server", "卡密、设备、账号、历史、远程端点维护、端点错误遥测", "授权运营、远程维护、设备与卡密管理"],
        ],
        widths=[1.2, 1.2, 2.4, 1.5],
        font_size=8.8,
    )
    callout(
        doc,
        "客户兼容优先",
        "AINexus 已面向客户分发。任何新增功能、数据库迁移、部署和远程维护能力，都必须先保证旧客户端、旧授权票据、旧 SQLite 数据、Docker volume 和本地配置可继续使用；远程能力不能在客户不知情时改变端点顺序、开关、凭证、模型或代理行为。",
    )


def install(doc):
    doc.add_heading("2. 安装、授权与接入", level=1)
    doc.add_heading("2.1 桌面客户端快速开始", level=2)
    numbers(
        doc,
        [
            "下载 AINexus 6.7.2 对应平台安装包。macOS 解压后移动到“应用程序”；Windows 解压后运行 AINexus.exe。",
            "首次启动输入在线卡密并联网激活。授权成功后，最近一次在线校验成功可离线宽限 30 天。",
            "在主界面添加或启用至少一个端点，建议先点击端点卡片“测试”。",
            "把 Codex CLI、Claude Code、OpenClaw 或 Hermes Agent 的 base URL 指向 AINexus。",
        ],
    )
    doc.add_heading("2.2 服务器模式快速开始", level=2)
    para(doc, "源码启动：go run ./cmd/server", bold=True)
    bullets(doc, ["API Provider：http://127.0.0.1:3000", "Web 管理界面：http://127.0.0.1:3000/ui/", "健康检查：http://127.0.0.1:3000/health", "Docker/NAS 场景应挂载数据目录并设置强 Basic Auth 密码。"])
    add_table(
        doc,
        ["环境变量", "说明", "默认/建议"],
        [
            ["AINEXUS_PORT", "HTTP 监听端口", "3000"],
            ["AINEXUS_LISTEN_MODE", "local 仅本机；lan 监听所有网卡", "默认 local；局域网共享时使用 lan"],
            ["AINEXUS_DATA_DIR", "数据目录", "~/.AINexus；容器内通常为 /data"],
            ["AINEXUS_DB_PATH", "SQLite 数据库路径", "<数据目录>/ainexus.db"],
            ["AINEXUS_BASIC_AUTH_ENABLED", "保护 Web UI 与 /api/ 管理接口", "true"],
            ["AINEXUS_BASIC_AUTH_USERNAME", "登录用户名", "admin"],
            ["AINEXUS_BASIC_AUTH_PASSWORD", "登录密码", "生产环境必须显式设置强密码"],
        ],
        widths=[2.0, 2.55, 1.75],
    )
    doc.add_heading("2.3 客户端接入示例", level=2)
    add_table(
        doc,
        ["客户端", "Base URL", "认证占位", "说明"],
        [
            ["Codex CLI", "http://127.0.0.1:3000/v1", "ainexus-local", "wire_api 使用 responses；上游认证由 AINexus 端点或 Token Pool 管理。"],
            ["Claude Code", "http://127.0.0.1:3000", "ANTHROPIC_AUTH_TOKEN=ainexus", "可在 settings.json 中设置 ANTHROPIC_BASE_URL。"],
            ["OpenAI SDK", "http://127.0.0.1:3000/v1", "任意本地占位 key", "实际上游模型和密钥由 AINexus 选择。"],
            ["局域网设备", "http://服务器IP:3000/v1", "按客户端要求填写", "仅在 listenMode=lan 且网络安全措施到位时使用。"],
        ],
        widths=[1.35, 2.2, 1.55, 1.2],
        font_size=8.7,
    )


def desktop(doc):
    doc.add_heading("3. 桌面客户端模式：每个界面与按钮", level=1)
    image(doc, "desktop-home.png", "图 1：AINexus 6.7.2 桌面客户端主界面。", width=6.25)
    add_table(
        doc,
        ["区域/按钮", "作用", "操作说明"],
        [
            ["修复 Codex 会话", "打开 Codex 会话可见性修复", "用于账号/API Key 切换后官方 Codex 会话侧边栏不可见的恢复。"],
            ["AI 助手", "打开内置 AI Agent", "通过当前启用端点回答配置检查、排障或本机 Agent 相关任务。"],
            ["智能体配置", "打开 Agent Provider 管理", "为 Claude Code、Codex CLI、OpenClaw、Hermes 等写入或还原 AINexus Provider。"],
            ["端口号", "打开访问设置", "修改端口、local/lan 访问模式，查看地址和活动连接。"],
            ["齿轮", "打开设置", "配置语言、主题、通知、关闭行为、代理、授权和故障转移。"],
            ["统计周期标签", "今日/昨日/本周/本月/历史", "历史会打开月度归档统计。"],
            ["统计筛选", "按端点、客户端 IP 或 IP 关键字筛选", "用于定位某端点或某设备的请求和 Token 消耗。"],
            ["端点列表收起", "折叠端点区域", "端点很多时便于查看统计或日志。"],
            ["视图按钮", "卡片视图/紧凑列表视图", "卡片看详情，列表适合大量端点快速扫描。"],
            ["筛选下拉：类型", "按转换器筛选端点", "Claude、Gemini、OpenAI、OpenAI2、DeepSeek、Kimi、Poe。"],
            ["筛选下拉：可用性", "按可用/未知/不可用筛选", "基于测试和运行时状态。"],
            ["筛选下拉：启用状态", "按启用/禁用筛选", "筛选开启时拖拽排序禁用。"],
            ["启动器", "打开终端启动器", "启动 Claude Code 或 Codex，并可选择项目目录和历史会话。"],
            ["数据同步", "打开备份与同步", "支持 WebDAV、本地目录、S3 兼容存储。"],
            ["添加端点", "打开新增端点弹窗", "新增 API Key、Token Pool、Codex Pool 或 Claude OAuth Pool 端点。"],
            ["端点开关", "启用/禁用端点", "禁用端点不会参与轮换。"],
            ["端点测试", "测试上游连接", "结果影响卡片图标和排障判断。"],
            ["端点复制", "克隆端点配置", "生成副本并打开新增端点弹窗。"],
            ["端点编辑", "编辑该端点", "进入端点配置弹窗。"],
            ["端点删除", "删除端点", "删除前二次确认；生产环境建议先备份数据库。"],
        ],
        widths=[1.75, 1.85, 2.7],
        font_size=8.4,
    )

    doc.add_heading("3.1 添加/编辑端点", level=2)
    image(doc, "desktop-endpoint-modal-api-key.png", "图 2：API Key 端点编辑窗口。", width=3.9)
    image(doc, "desktop-endpoint-modal-codex-pool.png", "图 3：Codex Token Pool 模式会锁定 Codex 上游和 OpenAI2 Responses 转换器。", width=3.9)
    add_table(
        doc,
        ["字段/按钮", "作用", "6.7.2 使用说明"],
        [
            ["名称", "端点显示名", "必须唯一；用于统计、远程快照、默认端点和错误遥测聚合。"],
            ["认证方式", "api_key / token_pool / codex_token_pool / claude_oauth_token_pool", "每个端点只能选择一种认证方式；Token Pool 模式隐藏 API Key 字段。"],
            ["管理 Token Pool", "保存端点并进入凭据池", "新增端点时会先保存，再打开凭据管理。"],
            ["API 地址", "上游基础 URL", "Codex Pool 和 Claude OAuth Pool 为固定地址，界面会锁定。"],
            ["代理设置", "端点级 HTTP/SOCKS5 代理", "优先于全局代理，适合 ChatGPT/Codex 或特定上游走代理。"],
            ["API 密钥/眼睛", "填写并显示/隐藏密钥", "不要在截图、远程维护或日志里暴露明文。"],
            ["转换器", "选择请求/响应协议转换", "OpenAI2 对应 Responses API，是 Codex CLI 推荐链路。"],
            ["模型 + 检测", "覆盖模型或从上游拉取模型列表", "模型字段为空时尽量尊重客户端原始模型；兼容上游需按转换器填写。"],
            ["推理", "端点级 thinking effort", "DeepSeek 支持上游默认/High/Max；其他支持 Low/Medium/High/XHigh。"],
            ["上游强制流式", "总是向上游发流式请求", "非流式客户端由 AINexus 聚合响应。"],
            ["快速模式", "Codex Pool 使用 service_tier=fast", "可能消耗更高 Codex 额度；远程维护也可开关。"],
            ["限制并发", "端点同时请求数上限", "0 表示不限制；达到上限时普通轮询请求跳到下一个端点。"],
            ["备注", "端点说明", "会显示在端点卡片和远程快照中。"],
            ["取消/保存", "放弃或保存配置", "保存会校验必填项、名称唯一性和并发格式。"],
        ],
        widths=[1.5, 2.2, 2.6],
        font_size=8.5,
    )

    doc.add_heading("3.2 Token Pool 管理", level=2)
    image(doc, "desktop-token-pool.png", "图 4：Codex Token Pool 管理窗口，含代理、导入、认证、额度和凭证列表。", width=6.25)
    add_table(
        doc,
        ["按钮/区域", "作用", "适用模式/说明"],
        [
            ["代理地址 + 保存/清空", "设置该 Pool 专用上游代理", "API Token Pool、Codex Pool、Claude OAuth Pool。"],
            ["导入 JSON", "粘贴单个对象、数组或 {items:[...]}", "用于批量导入 access_token / refresh_token / setup-token。"],
            ["覆盖已有账号/邮箱", "导入时更新旧账号", "适合续期或刷新同一账号凭据。"],
            ["导入", "提交文本框凭据", "显示新增、更新、跳过、失败数量。"],
            ["导入文件", "多文件批量导入 JSON", "适合账号多的 Codex Pool。"],
            ["发现 Claude", "发现本机 Claude OAuth token", "仅 Claude OAuth Token Pool。"],
            ["认证", "启动 ChatGPT 设备码登录", "仅 Codex Token Pool，会显示验证链接和用户代码。"],
            ["刷新", "重新加载凭据列表", "所有 Pool。"],
            ["刷新额度", "刷新 Codex 额度窗口", "仅 Codex Pool；6.7.x 会持久化额度快照。"],
            ["凭证启用/停用", "控制单条凭据是否参与轮换", "异常凭据建议先停用再排查。"],
            ["更多：查看错误", "显示该凭证最后错误", "便于区分 401、429、网络错误和上游错误。"],
            ["更多：刷新 token", "用 refresh_token 刷新 access_token", "6.7.x 强化单条凭证刷新。"],
            ["更多：更新 token", "手动粘贴新的 access_token", "用于外部刷新后回填。"],
            ["更多：用量", "查看凭证级请求、错误、输入/输出 Token", "6.7.x 重点能力，便于定位高消耗账号。"],
            ["更多：重置额度", "消耗 Codex reset credit 重置额度", "仅有可用重置次数时使用。"],
            ["更多：删除", "删除单条凭据", "删除前确认，无法参与后续轮换。"],
            ["关闭", "关闭管理窗口", "不影响代理运行。"],
        ],
        widths=[1.7, 2.25, 2.35],
        font_size=8.3,
    )

    doc.add_heading("3.3 设置、授权与访问", level=2)
    image(doc, "desktop-settings.png", "图 5：桌面设置窗口。", width=3.9)
    image(doc, "desktop-access-settings.png", "图 6：访问设置窗口。", width=5.1)
    image(doc, "desktop-license-gate.png", "图 7：启动授权门禁。", width=4.6)
    add_table(
        doc,
        ["设置项/按钮", "作用", "说明"],
        [
            ["语言", "切换中文/英文", "界面会按当前语言重新渲染。"],
            ["主题", "选择明亮、暗黑和扩展主题", "自动模式可按时间切换昼夜主题。"],
            ["自动模式", "启用自动主题", "开启后可配置白天/夜间主题。"],
            ["通知方式", "Claude Code 任务完成提醒", "修改后通常需重启终端生效。"],
            ["窗口关闭行为", "关闭/托盘/每次询问", "长期代理建议最小化到托盘。"],
            ["代理设置", "全局 HTTP/SOCKS5 代理", "为空直连；端点级代理优先。"],
            ["授权信息", "状态、到期时间、剩余天数、卡类型", "到期前续期避免代理服务受限。"],
            ["访问设置：端口", "修改监听端口", "保存后立即生效；现有流式连接可能短暂中断。"],
            ["访问设置：local/lan", "本机或局域网访问", "lan 模式需防火墙/VPN/强密码。"],
            ["局域网发现", "发现可添加的 AINexus", "可快速把同网段实例作为端点。"],
            ["使用中的连接", "查看活动连接分类、IP、路径、持续时间", "用于重启前确认是否有人正在使用。"],
            ["授权门禁：刷新", "重新校验授权状态", "网络恢复或续期后使用。"],
            ["授权门禁：激活", "提交卡密激活/续期", "成功后解除启动门禁。"],
        ],
        widths=[1.8, 2.1, 2.4],
        font_size=8.5,
    )

    doc.add_heading("3.4 数据同步、启动器、AI Agent 与 Provider", level=2)
    image(doc, "desktop-data-sync-webdav.png", "图 8：WebDAV 备份同步。", width=4.7)
    image(doc, "desktop-data-sync-s3.png", "图 9：S3 兼容存储备份。", width=4.7)
    image(doc, "desktop-launcher.png", "图 10：启动器。", width=4.25)
    image(doc, "desktop-ai-agent.png", "图 11：AI Agent 工作台。", width=6.1)
    image(doc, "desktop-agent-provider.png", "图 12：Agent Provider 管理。", width=5.6)
    add_table(
        doc,
        ["界面/按钮", "作用", "说明"],
        [
            ["WebDAV / 本地备份 / S3备份", "切换备份提供方", "三种方式都支持备份、恢复和管理。"],
            ["测试连接", "验证 WebDAV 或 S3 可用性", "首次配置建议先测试。"],
            ["保存配置", "保存备份参数", "凭据写入本地数据库。"],
            ["备份", "立即创建备份", "升级、批量修改端点或远程维护前建议先备份。"],
            ["备份管理器", "查看/恢复/删除备份", "恢复前会做冲突检测或确认。"],
            ["启动器：CLI 类型", "Claude Code / Codex", "影响启动命令、会话恢复和配置目标。"],
            ["启动器：选择终端", "选择本机终端程序", "macOS/Windows 按系统能力显示。"],
            ["启动器：添加目录", "登记项目目录", "点击项目可快速启动 CLI。"],
            ["启动器：恢复会话", "选择历史会话继续工作", "配合 Codex 会话修复使用。"],
            ["AI Agent：新对话", "开始新本地会话", "旧记录仍保存在本机。"],
            ["AI Agent：发送", "提交任务", "需要有可用端点。"],
            ["AI Agent：过程详情", "展开工具调用和端点信息", "用于排障 Agent 运行结果。"],
            ["Agent Provider：全选/清空", "批量选择目标客户端", "目标包括 Claude Code、Codex、OpenClaw、Hermes 等。"],
            ["Agent Provider：创建缺失配置", "缺文件时创建", "新机器初始化常用。"],
            ["Agent Provider：选择备份/还原备份", "恢复历史配置", "覆盖前会保留备份。"],
            ["Agent Provider：覆盖为 AINexus", "把选中客户端指向 AINexus", "会写入本地 provider 配置。"],
        ],
        widths=[1.75, 2.05, 2.5],
        font_size=8.3,
    )


def server(doc):
    doc.add_heading("4. 服务器模式：Web 管理界面与按钮", level=1)
    image(doc, "server-dashboard.png", "图 13：服务器 Web UI 仪表板。", width=6.25)
    add_table(
        doc,
        ["导航/按钮", "作用", "说明"],
        [
            ["仪表板", "查看请求、成功率、Token、网络、Agent Provider、活动端点", "服务器日常巡检第一屏。"],
            ["端点", "管理端点和 Token Pool", "支持添加、编辑、删除、克隆、启停、测试、切换、拖拽排序。"],
            ["统计", "查看每日/每周/每月统计", "支持端点和客户端 IP 筛选。"],
            ["测试", "选择端点运行测试", "用于部署或修改后验证。"],
            ["设置", "授权、网络访问、故障转移策略", "服务器模式核心配置入口。"],
            ["语言按钮", "切换中文/英文", "左下角国旗按钮。"],
            ["主题按钮", "明暗主题切换", "左下角月亮/太阳按钮。"],
            ["管理 Agent Provider", "打开 Provider 弹窗", "可为服务器上的客户端配置 AINexus Provider。"],
        ],
        widths=[1.45, 2.25, 2.6],
    )
    doc.add_heading("4.1 端点页面", level=2)
    image(doc, "server-endpoints.png", "图 14：服务器端点管理列表。", width=6.25)
    image(doc, "server-endpoint-modal.png", "图 15：服务器新增端点窗口。", width=4.15)
    add_table(
        doc,
        ["按钮/字段", "作用", "说明"],
        [
            ["添加端点", "打开新增端点弹窗", "字段与桌面端一致。"],
            ["拖拽手柄", "调整轮换优先级", "排序影响故障转移顺序。"],
            ["复制 URL", "复制 API 地址", "用于迁移和核对。"],
            ["切换", "设为默认端点", "只有启用且非当前默认端点时显示。"],
            ["测试", "测试上游连接", "结果记录在浏览器本地状态。"],
            ["令牌池管理", "打开 Token Pool 管理", "认证方式决定显示 API/Codex/Claude OAuth 工具。"],
            ["启用开关", "启用或禁用端点", "禁用后不参与轮换。"],
            ["编辑", "修改端点字段", "保留旧字段兼容。"],
            ["克隆", "复制端点为新记录", "适合相同上游不同模型。"],
            ["删除", "删除端点", "生产环境先备份。"],
            ["获取模型", "拉取上游模型列表", "需 API 地址和可用密钥/凭据。"],
            ["上游强制流式", "启用请求级强制流式", "适合只接受 SSE 的上游。"],
            ["限制并发", "该端点并发上限", "0 为不限制。"],
        ],
        widths=[1.55, 2.2, 2.55],
        font_size=8.5,
    )
    doc.add_heading("4.2 统计、测试与设置页面", level=2)
    image(doc, "server-stats.png", "图 16：服务器统计页面。", width=6.25)
    image(doc, "server-testing.png", "图 17：服务器测试页面。", width=6.25)
    image(doc, "server-settings.png", "图 18：服务器设置页面。", width=6.25)
    image(doc, "server-agent-provider-modal.png", "图 19：服务器 Web UI 的 Agent Provider 弹窗。", width=4.7)
    add_table(
        doc,
        ["页面/按钮", "作用", "说明"],
        [
            ["统计：每日/每周/每月", "切换统计周期", "统计总请求、成功、错误、输入/输出 Token。"],
            ["统计：全部端点/全部 IP/搜索 IP", "过滤统计", "用于定位某设备或端点消耗。"],
            ["统计：清除筛选", "恢复全量统计", "避免筛选状态造成误判。"],
            ["测试：选择端点", "选择一个启用端点", "无启用端点时会提示。"],
            ["测试：运行测试", "触发服务端测试", "展示延迟、响应或错误。"],
            ["设置：刷新状态", "刷新授权状态", "续期后确认票据状态。"],
            ["设置：激活/续期", "提交卡密", "服务器模式成功后通常需重启服务恢复代理。"],
            ["设置：访问模式", "local/lan", "lan 时同网段可访问代理，必须加强网络安全。"],
            ["设置：保存并生效", "保存网络访问", "立即生效。"],
            ["设置：恢复端点策略", "恢复端点如何回归轮换", "首选降权更稳妥，自动回切接近旧行为。"],
            ["设置：冷却秒数", "错误分类冷却", "额度不足、限流、上游错误、网络错误、Token 不可用、配置错误。"],
            ["Agent Provider：覆盖为 AINexus", "写入 provider 配置", "支持备份与还原。"],
        ],
        widths=[1.75, 2.05, 2.5],
        font_size=8.5,
    )


def license_admin(doc):
    doc.add_heading("5. 授权服务器后台与远程维护", level=1)
    image(doc, "license-admin-remote.png", "图 20：授权服务器后台的设备、远程端点维护、端点错误遥测与远程 Token Pool 管理。", width=6.25)
    add_table(
        doc,
        ["界面/按钮", "作用", "安全边界"],
        [
            ["设备 Tab", "查看设备、设备 ID/IP、版本、授权状态", "设备 ID 和 IP 默认应脱敏，必要时手动显示。"],
            ["卡密 Tab", "生成、禁用、筛选卡密", "服务器仅保存卡密哈希，不保存明文卡密。"],
            ["账号 Tab", "后台账号与权限管理", "权限包括卡密、设备、远程查看、远程维护、查看密钥等。"],
            ["审计历史 Tab", "查看后台操作历史", "远程写操作必须可追踪。"],
            ["刷新", "重新加载后台数据", "不改变客户端配置。"],
            ["退出", "退出后台会话", "共享机器上必须退出。"],
            ["设备：明细", "展开设备详情", "加载卡密兑换、远程状态和遥测。"],
            ["设备：备注", "维护客户/设备说明", "便于售后识别设备。"],
            ["设备：修改到期", "调整设备授权到期", "应记录审计。"],
            ["设备：禁用当前", "禁用当前激活", "客户客户端会在下次校验后失效。"],
            ["远程维护：新增端点", "下发 endpoint.create 命令", "客户端离线时命令会等待或超时；不得上传密钥明文到普通列表。"],
            ["远程维护：上移/下移", "调整远程端点顺序", "必须保留本地兜底，服务器不可用时本地配置继续生效。"],
            ["远程维护：改URL/改Key/改并发", "下发端点字段更新", "改 Key 属敏感操作，应限权、审计、避免明文落日志。"],
            ["远程维护：开启/关闭快速", "切换 Codex Fast Mode", "可能消耗更高 Codex 额度，需明确告知客户。"],
            ["远程维护：停用/启用", "切换远程端点开关", "默认不应在客户不知情时修改。"],
            ["远程维护：删除", "删除远程端点", "高风险操作，必须确认和审计。"],
            ["端点错误遥测", "查看近 24 小时/近 7 天错误分类", "只上传端点、模型、错误分类、状态码、时间窗口等聚合信息，不上传 prompt/response。"],
            ["Token Pool：停用/启用", "远程控制凭据是否参与轮换", "不暴露 access token 明文。"],
            ["Token Pool：改Token", "更新远程凭据 token", "应最小权限并避免日志记录。"],
            ["Token Pool：查看Token", "一次性查看明文", "需要 HTTPS 安全上下文和专门权限；结果应即用即弃。"],
        ],
        widths=[1.75, 2.15, 2.4],
        font_size=8.1,
    )
    callout(
        doc,
        "远程维护原则",
        "远程管理必须最小权限、可审计、可回滚。客户端应区分本地用户修改和服务器策略下发；发生冲突时默认保护客户当前可用配置。服务器不可用时，客户本地配置和旧授权宽限仍应继续发挥作用。",
    )


def reference(doc):
    doc.add_heading("6. API、转换器与运维参考", level=1)
    doc.add_heading("6.1 代理路径", level=2)
    add_table(
        doc,
        ["路径", "用途"],
        [
            ["/", "主代理路由，承接 Claude/OpenAI/Gemini 兼容请求并执行转换、轮换、故障转移。"],
            ["/v1/messages/count_tokens", "Token 计数。"],
            ["/v1/models", "模型列表，带缓存。"],
            ["/health", "健康检查。"],
            ["/stats", "统计数据。"],
            ["/ui/", "服务器模式 Web UI。"],
            ["/api/*", "服务器管理 API，受 Basic Auth 和 CORS 防护。"],
        ],
        widths=[2.15, 4.15],
    )
    doc.add_heading("6.2 认证模式与转换器", level=2)
    add_table(
        doc,
        ["认证模式", "说明", "典型用途"],
        [
            ["api_key", "标准 API Key 认证", "单个上游 API 密钥。"],
            ["token_pool", "普通 API Token 池", "多个 API Token 轮换、失败隔离、凭证级统计。"],
            ["codex_token_pool", "ChatGPT/Codex 登录凭据池", "Codex CLI、Responses API、额度快照、reset credits。"],
            ["claude_oauth_token_pool", "实验性 Claude OAuth Token Pool", "Claude Code 订阅 OAuth 凭据导入与发现。"],
        ],
        widths=[1.65, 2.45, 2.2],
    )
    add_table(
        doc,
        ["转换器", "上游协议", "模型建议"],
        [
            ["claude", "Claude / Anthropic", "可留空或覆盖模型。"],
            ["openai", "OpenAI Chat Completions", "通常必填。"],
            ["openai2", "OpenAI Responses", "Codex CLI 推荐，通常必填。"],
            ["gemini", "Google Gemini", "通常必填。"],
            ["deepseek", "DeepSeek OpenAI Chat 兼容", "必填，thinking 可使用上游默认/High/Max。"],
            ["kimi", "Kimi/Moonshot OpenAI Chat 兼容", "必填。"],
            ["poe", "Poe bot OpenAI Chat 兼容", "按 bot 名填写。"],
        ],
        widths=[1.35, 2.6, 2.35],
    )
    doc.add_heading("6.3 故障转移与错误分类", level=2)
    add_table(
        doc,
        ["错误类别", "默认冷却含义", "排查建议"],
        [
            ["quota_exhausted", "额度不足冷却", "刷新额度、切换账号、检查 reset credits。"],
            ["rate_limited", "限流冷却", "降低并发、延长冷却、增加备用端点。"],
            ["upstream_error", "上游 5xx 或协议异常", "检查上游状态和转换器。"],
            ["network_error", "网络连接失败", "检查代理、DNS、防火墙。"],
            ["token_unavailable", "Token 不可用", "刷新 token、停用异常凭据。"],
            ["config_error", "配置错误", "检查 API URL、模型、认证模式和密钥。"],
        ],
        widths=[1.55, 2.2, 2.55],
    )
    doc.add_heading("6.4 升级与备份检查清单", level=2)
    bullets(
        doc,
        [
            "升级前备份 ~/.AINexus/ainexus.db、服务器数据目录或 Docker volume。",
            "确认旧在线票据、离线宽限、卡密兑换记录仍可校验。",
            "SQLite 迁移必须幂等，不删除旧字段或旧表；重命名/拆表要保留旧读路径。",
            "默认端口、授权服务器地址、代理行为和端点认证模式不要随意改变。",
            "远程维护默认不能暴露客户 API Key、refresh token、access token 明文。",
            "升级安装包应来自可信发布渠道；升级失败时先恢复数据库或备份目录，再回退到上一版可用程序。",
        ],
    )


def main():
    doc = Document()
    configure(doc)
    cover(doc)
    overview(doc)
    install(doc)
    desktop(doc)
    server(doc)
    license_admin(doc)
    reference(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
